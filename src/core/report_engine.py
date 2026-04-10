"""ReportEngine — Excel/CSV/Word/PDF export (§5A2, §5D1)."""

from __future__ import annotations

import csv
import logging
from datetime import datetime
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


class ReportEngine:
    """Generate reports in Excel, CSV, Word, and PDF formats."""

    def __init__(self, output_dir: str = "data/exports") -> None:
        self._output_dir = Path(output_dir)
        self._output_dir.mkdir(parents=True, exist_ok=True)

    def export_test_cases_excel(
        self, cases: list[dict[str, str]], chip_name: str
    ) -> str:
        """Export test cases to Excel file."""
        try:
            import openpyxl
            from openpyxl.styles import Font
        except ImportError:
            logger.error("openpyxl not installed")
            return ""

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Test Cases"

        headers = ["Test Item", "Parameter", "Condition", "Expected Value", "Test Method", "Priority"]
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.font = Font(bold=True)

        for row_idx, case in enumerate(cases, 2):
            ws.cell(row=row_idx, column=1, value=case.get("test_item", ""))
            ws.cell(row=row_idx, column=2, value=case.get("parameter", ""))
            ws.cell(row=row_idx, column=3, value=case.get("condition", ""))
            ws.cell(row=row_idx, column=4, value=case.get("expected_value", ""))
            ws.cell(row=row_idx, column=5, value=case.get("test_method", ""))
            ws.cell(row=row_idx, column=6, value=case.get("priority", ""))

        # Auto-width
        for col in ws.columns:
            max_len = max(len(str(cell.value or "")) for cell in col)
            ws.column_dimensions[col[0].column_letter].width = min(max_len + 2, 50)

        ws.freeze_panes = "A2"

        ts = datetime.now().strftime("%Y%m%d")
        filename = f"{chip_name}_test_cases_{ts}.xlsx"
        path = self._output_dir / filename
        wb.save(str(path))
        return str(path)

    def export_test_cases_csv(
        self, cases: list[dict[str, str]], chip_name: str
    ) -> str:
        """Export test cases to CSV file."""
        ts = datetime.now().strftime("%Y%m%d")
        filename = f"{chip_name}_test_cases_{ts}.csv"
        path = self._output_dir / filename

        with open(path, "w", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(
                f,
                fieldnames=["test_item", "parameter", "condition", "expected_value", "test_method", "priority"],
            )
            writer.writeheader()
            writer.writerows(cases)

        return str(path)

    def generate_word(
        self, data: dict[str, Any], title: str = "Report"
    ) -> str:
        """Generate a Word report."""
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed")
            return ""

        doc = Document()
        doc.add_heading(title, level=0)

        if "summary" in data:
            doc.add_paragraph(data["summary"])

        if "comparison_table" in data:
            table_data = data["comparison_table"]
            if table_data:
                chips = list(next(iter(table_data.values())).keys()) if table_data else []
                table = doc.add_table(rows=1 + len(table_data), cols=1 + len(chips))
                table.style = "Table Grid"
                table.cell(0, 0).text = "Parameter"
                for i, chip in enumerate(chips):
                    table.cell(0, i + 1).text = chip
                for row_idx, (param, values) in enumerate(table_data.items(), 1):
                    table.cell(row_idx, 0).text = param
                    for col_idx, chip in enumerate(chips):
                        v = values.get(chip, {})
                        table.cell(row_idx, col_idx + 1).text = str(v.get("typ", "N/A")) if isinstance(v, dict) else str(v or "N/A")

        if "analysis" in data:
            doc.add_heading("Analysis", level=1)
            doc.add_paragraph(data["analysis"])

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"report_{ts}.docx"
        path = self._output_dir / filename
        doc.save(str(path))
        return str(path)

    def generate_pdf(
        self, data: dict[str, Any], title: str = "Report"
    ) -> str:
        """Generate a PDF report (using reportlab)."""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
        except ImportError:
            logger.error("reportlab not installed")
            return ""

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"report_{ts}.pdf"
        path = self._output_dir / filename

        doc = SimpleDocTemplate(str(path), pagesize=A4)
        styles = getSampleStyleSheet()
        story = [
            Paragraph(title, styles["Title"]),
            Spacer(1, 12),
        ]

        if "summary" in data:
            story.append(Paragraph(data["summary"], styles["Normal"]))
            story.append(Spacer(1, 12))

        if "analysis" in data:
            story.append(Paragraph("Analysis", styles["Heading1"]))
            story.append(Paragraph(data["analysis"], styles["Normal"]))

        doc.build(story)
        return str(path)

    def generate_excel(
        self, data: dict[str, Any], title: str = "Report"
    ) -> str:
        """Generate an Excel report with data tables."""
        try:
            import openpyxl
            from openpyxl.styles import Font
        except ImportError:
            return ""

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = title[:31]
        ws.cell(row=1, column=1, value=title).font = Font(bold=True, size=14)

        row = 3
        if "comparison_table" in data:
            for param, values in data["comparison_table"].items():
                ws.cell(row=row, column=1, value=param)
                for col, (chip, val) in enumerate(values.items(), 2):
                    if row == 3:
                        ws.cell(row=2, column=col, value=chip).font = Font(bold=True)
                    ws.cell(row=row, column=col, value=str(val.get("typ", "")) if isinstance(val, dict) else str(val or ""))
                row += 1

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"report_{ts}.xlsx"
        path = self._output_dir / filename
        wb.save(str(path))
        return str(path)
